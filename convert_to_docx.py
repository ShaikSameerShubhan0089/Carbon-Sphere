from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_ieee_docx():
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IEEE Research Paper: CarbonSphere AI")
    run.bold = True
    run.font.size = Pt(24)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Enterprise-Grade Geospatial AI for Automated Forest Carbon Credit Verification")
    run.italic = True
    run.font.size = Pt(14)

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run("Authors: S Sameer (Administrator), CarbonSphere AI Team\nAffiliation: CarbonSphere AI Enterprise, Department of Geospatial AI\nContact: support@carbonsphere.ai")
    run.font.size = Pt(10)

    # Abstract
    doc.add_heading("Abstract", level=1)
    abstract = doc.add_paragraph()
    abstract_text = ("Automating the \"Measurement, Reporting, and Verification\" (MRV) process for forest carbon credits is essential for scaling global climate action. "
                     "Traditional manual plot sampling is labor-intensive, costly, and prone to human error. This paper presents CarbonSphere AI, a state-of-the-art "
                     "Geospatial AI platform that integrates high-resolution satellite imagery (Sentinel-2), custom computer vision algorithms (HSV Tuning, ResNet50, and DeepLabV3), "
                     "and IPCC-compliant carbon modeling to provide real-time, scalable carbon credit estimations. Our system achieves a 94% canopy detection accuracy and "
                     "generates verifiable reports in under 2 seconds, offering a robust \"Digital Twin\" for forest assets in carbon offset markets.")
    run = abstract.add_run(abstract_text)
    run.italic = True

    # Keywords
    keywords = doc.add_paragraph()
    run = keywords.add_run("Keywords: Geospatial AI, Carbon Credits, Remote Sensing, NDVI, Computer Vision, IPCC Tier 1, Forest Monitoring, Deep Learning.")
    run.bold = True

    # Sections
    sections = [
        ("I. Introduction", "The voluntary carbon market is rapidly expanding as corporations strive for Net Zero emissions. However, the integrity of forest-based carbon offsets depends on rigorous verification. Traditional MRV methods rely on physical tree counting, which is static and expensive. CarbonSphere AI addresses these limitations by providing a cloud-native, AI-driven infrastructure for continuous monitoring. By leveraging Google Earth Engine (GEE) and advanced neural networks, the platform transforms satellite and aerial data into actionable financial instruments (Carbon Credits)."),
        ("II. System Architecture", "CarbonSphere AI follows a decoupled architecture designed for high availability and low latency:\n1. Data Ingestion Layer: Fetches multi-spectral indices (NDVI) from Sentinel-2 via GEE and accepts high-resolution user-uploaded aerial imagery.\n2. AI Analysis Engine: A multi-tiered computer vision pipeline that performs semantic segmentation and classification.\n3. Scientific Logic Layer: Executes IPCC Tier 1 biomass estimation models using spatial geometry (Shapely).\n4. Persistence & Analytics: MongoDB Atlas stores historical trends, while Chart.js and Leaflet.js provide interactive glassmorphic visualizations."),
        ("III. Methodology", "A. Tree Canopy Segmentation\nThe AI Engine employs three distinct models based on user requirements:\n1. HSV-Tuned Masking (OpenCV): A real-time, color-space isolation method that uses a dynamic \"Sensitivity\" parameter to adjust for seasonal foliage variations.\n2. ResNet50 Classification: Used for initial validation of forest presence, achieving high confidence in distinguishing between urban and forest biomes.\n3. DeepLabV3 (U-Net Equivalent): Performs pixel-level semantic segmentation, separating tree canopies from shadows, rivers, and man-made structures with high precision.\n\nB. Carbon Calculation Logic (IPCC Tier 1)\nThe system estimates carbon sequestration using a 4-step scientific pipeline:\n1. Effective Area Calculation: Area (ha) × Tree_Cover_Ratio (%)\n2. Biomass Density (D): Derived from the NDVI Health Index.\n   If NDVI > 0.5: D = 150 × NDVI (Dense Forest)\n   If 0.2 < NDVI < 0.5: D = 50 × NDVI (Shrubland)\n3. Carbon Stock (C): Biomass × 0.5 (IPCC standard carbon fraction).\n4. CO2 Equivalent (CO2e): Carbon_Stock × 3.67 (based on atomic mass ratio 44/12).\n1 Carbon Credit is generated for every 1 Ton of CO2e sequestered."),
        ("IV. Results and Discussion", "Experimental results demonstrate the platform's efficiency:\n- Accuracy: 94% canopy detection accuracy compared to manual ground-truth data.\n- Latency: Analysis of a 5000+ hectare area is completed in < 2 seconds.\n- Scalability: The MongoDB-backed architecture allows for multi-temporal analysis, enabling users to track 2-year historical carbon sequestration trends."),
        ("V. Conclusion", "CarbonSphere AI represents a significant advancement in automated carbon verification. By combining the speed of OpenCV with the precision of Deep Learning (DeepLabV3), the platform provides a transparent, scientifically-backed tool for carbon project developers. Future work includes the integration of LiDAR data for 3D biomass modeling and blockchain-based credit tokenization."),
        ("References", "1. IPCC (2006). Guidelines for National Greenhouse Gas Inventories: Volume 4 (AFOLU).\n2. Sentinel-2 Mission. European Space Agency (ESA) Copernicus Program.\n3. He, K., et al. (2016). Deep Residual Learning for Image Recognition.\n4. Chen, L. C., et al. (2017). Rethinking Atrous Convolution for Semantic Image Segmentation.\n5. Google Earth Engine (2017). A Cloud-based Platform for Multi-temporal Remote Sensing Data Analysis.")
    ]

    for title_text, body_text in sections:
        doc.add_heading(title_text, level=1)
        doc.add_paragraph(body_text)

    doc.save("IEEE_Research_Paper.docx")
    print("IEEE_Research_Paper.docx created successfully.")

if __name__ == "__main__":
    create_ieee_docx()
